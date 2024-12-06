from docx import Document
import pandas as pd
import json

# 워드 파일 오픈
doc = Document("")

header = ["Global Variable","Input Parameters", "Output Parameters"]

VALUE = {}


def Append_Variable(name, type, description):
    if name in VALUE:
        # 키가 존재
        if VALUE[name][0] == type and VALUE[name][1] == description:
            return
        # else:
        #     VALUE[name].append([type, description])

    VALUE[name] = [type, description]
    

for para in doc.paragraphs:

    #para.text가 header에 포함된다면
    if para.text.strip() in header:
        # print(f"Header Found: {para.text.strip()}")
        
        # 현재 단락 바로 다음에 있는 표 출력
         # 헤더 바로 다음의 요소가 표인지 확인
        next_element = para._p.getnext()
        if next_element is not None and next_element.tag.endswith('tbl'):  # 표인지 확인
            for table in doc.tables:
                

                if table._tbl == next_element:

                    # print("Table Contents:")
                    for idx, row in enumerate(table.rows, start=1):

                        #제목 행 스킵
                        if idx == 1:
                            continue

                        cell_data = [cell.text.strip() for cell in row.cells]
                        
                        if cell_data[1] == "N/A" or cell_data[0] == "":
                            continue
                        
                        
                        # 글로벌 변수 추출
                        if para.text.strip() == header[0]:
                            Append_Variable(cell_data[0].replace(' ',''), cell_data[1], cell_data[3])

                        # Input / Output 변수 추출
                        else:
                            Append_Variable(cell_data[1].replace(' ',''), cell_data[2], cell_data[4])  

# for key, data in VALUE.items():
#     if len(data) > 1:
#         print(key, ":", data)

def save_as_excel():
    # 데이터 준비
    rows = []
    for key, data in VALUE.items():
        rows.append([key, data[0], data[1]])

    # DataFrame으로 변환
    df = pd.DataFrame(rows, columns=["Variable", "Type", "Description"])

    # 엑셀로 저장
    df.to_excel("output.xlsx", index=False)

    print("엑셀 파일이 생성되었습니다.")

def save_as_json():
    with open('value_data.json', 'w') as json_file:
        json.dump(VALUE, json_file, indent=4)
    
    print("데이터가 value_data.json 파일에 저장되었습니다.")
save_as_excel()
save_as_json()